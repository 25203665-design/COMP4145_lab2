from docx import Document
import pandas as pd

# Load the dataset
df = pd.read_csv('../../../../Downloads/Titanic.csv')

# Create a Word document
doc = Document()
doc.add_heading('Titanic Dataset Data Cleaning and Transformation Report', 0)

# Original Data Section
doc.add_heading('1. Original Dataset Overview', level=1)
doc.add_paragraph('The Titanic dataset contains information about passengers on the Titanic ship. Below is a preview of the first 10 rows:')

# Add a table for original data
table = doc.add_table(rows=1, cols=len(df.columns))
hdr_cells = table.rows[0].cells
for i, col in enumerate(df.columns):
    hdr_cells[i].text = col

for _, row in df.head(10).iterrows():
    row_cells = table.add_row().cells
    for i, val in enumerate(row):
        row_cells[i].text = str(val)

# Data Cleaning Section
doc.add_heading('2. Data Cleaning: Handling Missing Values in Age', level=1)
doc.add_paragraph('Issue Identified: The Age column has missing values (NaN). This can lead to biased analysis or errors in data mining algorithms that require complete data.')
doc.add_paragraph('Solution: Impute missing Age values with the median age to preserve the distribution and avoid outliers.')
median_age = df['Age'].median()
doc.add_paragraph(f'Median Age: {median_age:.1f}')

# Apply cleaning
df['Age'] = pd.to_numeric(df['Age'], errors='coerce')
df['Age'] = df['Age'].fillna(median_age)

# Show cleaned data
doc.add_paragraph('After imputation (showing PassengerId, Name, Age):')
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'PassengerId'
hdr_cells[1].text = 'Name'
hdr_cells[2].text = 'Age'

for _, row in df[['PassengerId', 'Name', 'Age']].head(10).iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(row['PassengerId'])
    row_cells[1].text = str(row['Name'])
    row_cells[2].text = str(row['Age'])

# Data Transformation Section
doc.add_heading('3. Data Transformation: Creating FamilySize Feature', level=1)
doc.add_paragraph('Purpose: To create a new feature that represents the total family size, which can be useful for survival analysis.')
doc.add_paragraph('Transformation: FamilySize = SibSp + Parch + 1 (including the passenger themselves).')
doc.add_paragraph('Importance: This feature can help identify patterns related to family groups, potentially improving model performance in predicting survival.')

# Apply transformation
df['FamilySize'] = df['SibSp'] + df['Parch'] + 1

# Show transformed data
doc.add_paragraph('After transformation (showing PassengerId, Name, SibSp, Parch, FamilySize):')
table = doc.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'PassengerId'
hdr_cells[1].text = 'Name'
hdr_cells[2].text = 'SibSp'
hdr_cells[3].text = 'Parch'
hdr_cells[4].text = 'FamilySize'

for _, row in df[['PassengerId', 'Name', 'SibSp', 'Parch', 'FamilySize']].head(10).iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(row['PassengerId'])
    row_cells[1].text = str(row['Name'])
    row_cells[2].text = str(row['SibSp'])
    row_cells[3].text = str(row['Parch'])
    row_cells[4].text = str(row['FamilySize'])

# Summary
doc.add_heading('4. Summary', level=1)
doc.add_paragraph('Data Cleaning: Missing Age values were imputed with the median (28.0) to ensure completeness without introducing bias.')
doc.add_paragraph('Data Transformation: A new FamilySize feature was added by combining SibSp and Parch, providing insights into family structures.')
doc.add_paragraph('These steps improve data quality for better data mining performance by handling missing data and creating meaningful features.')

# Save the document
doc.save('Titanic_Report.docx')
print("Word report created: Titanic_Report.docx")
